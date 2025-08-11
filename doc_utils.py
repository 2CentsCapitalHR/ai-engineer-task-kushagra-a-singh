import datetime
import json
import logging
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.text import WD_BREAK, WD_COLOR_INDEX
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

# Enhanced document structure patterns
DOCUMENT_PATTERNS = {
    "Articles of Association (AoA)": [
        r"articles?\s+of\s+association",
        r"constitutional\s+document",
        r"company\s+constitution",
        r"share\s+capital",
        r"director\s+power",
        r"shareholder\s+right",
    ],
    "Memorandum of Association (MoA/MoU)": [
        r"memorandum\s+of\s+(association|understanding)",
        r"company\s+object",
        r"registered\s+office",
        r"liability\s+of\s+member",
        r"initial\s+member",
    ],
    "Board Resolution (for Incorporation)": [
        r"board\s+resolution",
        r"directors?\s+resolution",
        r"resolved\s+that",
        r"board\s+meeting",
        r"quorum\s+present",
    ],
    "Shareholder Resolution (for Incorporation)": [
        r"shareholder\s+resolution",
        r"members?\s+resolution",
        r"special\s+resolution",
        r"ordinary\s+resolution",
        r"AGM\s+resolution",
    ],
    "Incorporation Application Form": [
        r"incorporation\s+application",
        r"company\s+registration",
        r"proposed\s+company\s+name",
        r"nature\s+of\s+business",
        r"authorized\s+share\s+capital",
    ],
    "Ultimate Beneficial Owner (UBO) Declaration Form": [
        r"UBO\s+declaration",
        r"ultimate\s+beneficial\s+owner",
        r"beneficial\s+ownership",
        r"ownership\s+structure",
        r"control\s+structure",
    ],
    "Standard Employment Contract (2024 update)": [
        r"employment\s+contract",
        r"employment\s+agreement",
        r"terms\s+of\s+employment",
        r"probation\s+period",
        r"notice\s+period",
        r"end\s+of\s+service",
    ],
    "Shareholder Agreement (SHA)": [
        r"shareholder\s+agreement",
        r"investment\s+agreement",
        r"transfer\s+restriction",
        r"tag\s+along",
        r"drag\s+along",
        r"pre-emption",
    ],
    "Non-Disclosure Agreement (NDA)": [
        r"non.?disclosure\s+agreement",
        r"confidentiality\s+agreement",
        r"confidential\s+information",
        r"proprietary\s+information",
        r"trade\s+secret",
    ],
    "Risk Policy Statement": [
        r"risk\s+policy",
        r"risk\s+management",
        r"risk\s+framework",
        r"risk\s+appetite",
        r"risk\s+tolerance",
    ],
    "FSRA License Application": [
        r"FSRA\s+license",
        r"financial\s+services\s+license",
        r"regulatory\s+approval",
        r"license\s+application",
        r"financial\s+services\s+activity",
    ],
}

# Document structure requirements
REQUIRED_SECTIONS = {
    "Articles of Association (AoA)": [
        "company name",
        "registered office",
        "objects",
        "share capital",
        "directors",
        "shareholders",
        "meetings",
        "accounts",
        "winding up",
    ],
    "Memorandum of Association (MoA/MoU)": [
        "company name",
        "registered office",
        "objects",
        "liability",
        "initial members",
        "share capital",
    ],
    "Board Resolution (for Incorporation)": [
        "date",
        "directors present",
        "quorum",
        "resolutions",
        "signatures",
    ],
    "Employment Contract": [
        "parties",
        "job title",
        "duties",
        "salary",
        "working hours",
        "probation",
        "notice period",
        "termination",
        "governing law",
    ],
    "Shareholder Agreement (SHA)": [
        "parties",
        "definitions",
        "share transfer",
        "management",
        "board composition",
        "dividend policy",
        "exit provisions",
    ],
}

# Signature detection patterns
SIGNATURE_PATTERNS = [
    r"signed\s+by",
    r"signature\s*:",
    r"authorized\s+signatory",
    r"director\s+signature",
    r"witness\s+signature",
    r"executed\s+by",
    r"in\s+witness\s+whereof",
    r"________________________",
    r"____________________",
]

# Date detection patterns
DATE_PATTERNS = [
    r"\d{1,2}[/-]\d{1,2}[/-]\d{4}",
    r"\d{1,2}\s+(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}",
    r"(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}",
    r"\d{1,2}(st|nd|rd|th)\s+(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}",
]


def extract_text_sections(docx_file) -> List[Dict[str, Any]]:
    """Enhanced text extraction with comprehensive structure detection"""
    try:
        doc = Document(docx_file)
        sections = []

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            section_info = {
                "index": i,
                "text": text,
                "raw": text,
                "style": para.style.name if para.style else "Normal",
                "type": "content",
            }

            # Enhanced clause/section detection
            clause_patterns = [
                r"^(Clause|Section|Article|Paragraph|Para)?\s*([\d\.]+)\s*[:\-]?\s*(.*)$",
                r"^(\d+\.?)\s+(.+)$",
                r"^([A-Z]\.?)\s+(.+)$",
            ]

            for pattern in clause_patterns:
                match = re.match(pattern, text, re.I)
                if match:
                    section_info.update(
                        {
                            "clause": (
                                match.group(2)
                                if len(match.groups()) >= 2
                                else match.group(1)
                            ),
                            "content": (
                                match.group(3)
                                if len(match.groups()) >= 3
                                else match.group(2)
                            ),
                            "type": "clause",
                        }
                    )
                    break

            # Detect headers
            if (
                (para.style and para.style.name.startswith("Heading"))
                or (len(text) < 100 and text.isupper())
                or (len(text.split()) <= 10 and text.endswith(":"))
            ):
                section_info["type"] = "header"

            # Detect signatures
            if any(re.search(pattern, text.lower()) for pattern in SIGNATURE_PATTERNS):
                section_info["type"] = "signature"

            # Detect dates
            if any(re.search(pattern, text) for pattern in DATE_PATTERNS):
                section_info["has_date"] = True

            # Extract defined terms
            defined_terms = re.findall(r'"([^"]+)"', text)
            if defined_terms:
                section_info["defined_terms"] = defined_terms

            sections.append(section_info)

        # Process tables if present
        for table_idx, table in enumerate(doc.tables):
            table_content = []
            for row in table.rows:
                row_content = [cell.text.strip() for cell in row.cells]
                if any(row_content):  # Skip empty rows
                    table_content.append(row_content)

            if table_content:
                sections.append(
                    {
                        "index": len(sections),
                        "text": f"Table {table_idx + 1}: "
                        + "; ".join([" | ".join(row) for row in table_content[:3]]),
                        "raw": str(table_content),
                        "type": "table",
                        "table_data": table_content,
                    }
                )

        return sections

    except Exception as e:
        logger.error(f"Error extracting text sections: {e}")
        return []


def detect_document_type(
    text: str, all_types: List[str], use_ai_fallback: bool = True
) -> str:
    """Enhanced document type detection with comprehensive pattern matching"""
    if not text:
        return "Unknown"

    text_lower = text.lower()

    # Direct type matching
    for doc_type in all_types:
        if doc_type.lower() in text_lower:
            return doc_type

    # Pattern-based detection with scoring
    type_scores = {}

    for doc_type, patterns in DOCUMENT_PATTERNS.items():
        score = 0
        for pattern in patterns:
            matches = len(re.findall(pattern, text_lower))
            score += matches * 2  # Weight multiple matches higher

        if score > 0:
            type_scores[doc_type] = score

    # Return highest scoring type if confidence is sufficient
    if type_scores:
        best_match = max(type_scores.items(), key=lambda x: x[1])
        if best_match[1] >= 2:  # Minimum confidence threshold
            return best_match[0]

    # Keyword-based fallback detection
    keyword_mapping = {
        "articles": "Articles of Association (AoA)",
        "memorandum": "Memorandum of Association (MoA/MoU)",
        "resolution": "Board Resolution (for Incorporation)",
        "employment": "Standard Employment Contract (2024 update)",
        "shareholder agreement": "Shareholder Agreement (SHA)",
        "NDA": "Non-Disclosure Agreement (NDA)",
        "UBO": "Ultimate Beneficial Owner (UBO) Declaration Form",
        "risk policy": "Risk Policy Statement",
        "license": "License Application Form",
    }

    for keyword, doc_type in keyword_mapping.items():
        if keyword.lower() in text_lower:
            return doc_type

    # AI fallback if enabled and available
    if use_ai_fallback:
        try:
            from adgm_rag import gemini_legal_analysis

            prompt = f"Identify the document type from this text (respond with just the document type): {text[:500]}"
            ai_result = gemini_legal_analysis(prompt)
            if isinstance(ai_result, dict) and ai_result.get("red_flag"):
                return ai_result["red_flag"]
        except Exception as e:
            logger.warning(f"AI fallback failed: {e}")

    return "Unknown"


def validate_document_structure(doc_type: str, metadata: Dict[str, Any]) -> List[str]:
    """Enhanced document structure validation with comprehensive checks"""
    issues = []

    # General validation rules
    if metadata.get("paragraph_count", 0) < 5:
        issues.append("Document appears too short for a legal document")

    # Type-specific validation
    required_sections = REQUIRED_SECTIONS.get(doc_type, [])

    if required_sections:
        # Check for presence of required sections
        content_lower = " ".join(
            [h.lower() for h in metadata.get("headers", [])]
        ).lower()
        missing_sections = []

        for section in required_sections:
            if section.lower() not in content_lower:
                missing_sections.append(section)

        if missing_sections:
            issues.append(
                f"Missing required sections: {', '.join(missing_sections[:3])}"
            )

    # Signature validation
    if doc_type in [
        "Articles of Association (AoA)",
        "Memorandum of Association (MoA/MoU)",
        "Board Resolution (for Incorporation)",
        "Shareholder Resolution (for Incorporation)",
    ]:
        if not metadata.get("signatures"):
            issues.append("Missing signature section")

    # Date validation for resolutions
    if "Resolution" in doc_type and not metadata.get("dates"):
        issues.append("Resolution missing date")

    # ADGM-specific validations
    if doc_type in [
        "Articles of Association (AoA)",
        "Memorandum of Association (MoA/MoU)",
    ]:
        # Check for ADGM jurisdiction
        all_text = " ".join(metadata.get("headers", []))
        if (
            "adgm" not in all_text.lower()
            and "abu dhabi global market" not in all_text.lower()
        ):
            issues.append("Document should reference ADGM jurisdiction")

    return issues


def add_comment_to_paragraph(
    paragraph,
    comment_text: str,
    severity: str = "Medium",
    citation: str = "",
    suggestion: str = "",
) -> None:
    """Enhanced commenting with comprehensive formatting and legal citations"""
    try:
        # Create comprehensive comment
        comment_parts = []

        if comment_text:
            comment_parts.append(f"ISSUE: {comment_text}")

        if citation:
            comment_parts.append(f"CITATION: {citation}")

        if suggestion:
            comment_parts.append(f"SUGGESTION: {suggestion}")

        full_comment = " | ".join(comment_parts)

        # Add a line break before the comment to improve readability
        paragraph.add_run().add_break()

        # Add comment run with severity-based formatting
        comment_run = paragraph.add_run(f"[ADGM REVIEW: {full_comment}]")

        # Enhanced color coding and formatting
        severity_config = {
            "High": {
                "color": RGBColor(220, 20, 60),
                "bold": True,
                "italic": True,
            },  # Crimson
            "Medium": {
                "color": RGBColor(255, 140, 0),
                "bold": False,
                "italic": True,
            },  # Dark Orange
            "Low": {
                "color": RGBColor(218, 165, 32),
                "bold": False,
                "italic": True,
            },  # Golden Rod
        }

        config = severity_config.get(severity, severity_config["Medium"])
        comment_run.font.color.rgb = config["color"]
        comment_run.font.bold = config["bold"]
        comment_run.font.italic = config["italic"]
        comment_run.font.size = Pt(9)  # Smaller, readable font size

        # Add highlighting
        if severity == "High":
            comment_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    except Exception as e:
        logger.error(f"Error adding comment to paragraph: {e}")


def add_hyperlink(
    paragraph,
    url: str,
    text: str,
    color: RGBColor = RGBColor(0, 102, 204),
    underline: bool = True,
):
    """Add a clickable hyperlink to a paragraph."""
    try:
        # Create the relationship id
        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        # Create a w:r element
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        # Set color - fix the RGBColor.rgb issue
        color_elem = OxmlElement("w:color")
        # RGBColor object has a hex property directly, not .rgb.hex
        try:
            hex_value = (
                color.hex
                if hasattr(color, "hex")
                else f"{color.rgb[0]:02x}{color.rgb[1]:02x}{color.rgb[2]:02x}"
            )
        except:
            hex_value = "0000ff"  # Default blue if all else fails
        color_elem.set(qn("w:val"), hex_value)
        rPr.append(color_elem)

        # Underline
        if underline:
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rPr.append(u)

        new_run.append(rPr)
        # Add text element
        text_elem = OxmlElement("w:t")
        text_elem.text = text
        new_run.append(text_elem)
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)
        return paragraph
    except Exception as e:
        logger.warning(f"Failed to add hyperlink '{text}': {e}")
        # Fallback to plain text
        fallback_run = paragraph.add_run(text)
        fallback_run.font.color.rgb = color
        return paragraph


def generate_dynamic_helpful_links(
    flagged_sections: List[Dict[str, Any]], doc_type: str
) -> List[Tuple[str, str]]:
    """Generate dynamic helpful resources based on actual issues found"""

    # Collect all categories and specific issues
    categories = set()
    specific_issues = set()

    for section in flagged_sections:
        category = section.get("category", "").lower()
        if category:
            categories.add(category)

        red_flag = section.get("red_flag", "").lower()
        if red_flag:
            # Extract key terms from red flags
            if "jurisdiction" in red_flag or "adgm" in red_flag:
                categories.add("jurisdiction")
            if "missing" in red_flag or "clause" in red_flag:
                categories.add("missing_clauses")
            if "format" in red_flag or "structure" in red_flag:
                categories.add("formatting")
            if "compliance" in red_flag or "regulation" in red_flag:
                categories.add("compliance")
            if "ambiguous" in red_flag or "unclear" in red_flag:
                categories.add("ambiguity")
            if "outdated" in red_flag or "old" in red_flag:
                categories.add("outdated")

    # Base resources map with more comprehensive links
    resources_map = {
        "jurisdiction": [
            (
                "ADGM Companies Regulations 2020",
                "https://en.adgm.thomsonreuters.com/rulebook/1-companies-regulations-2020",
            ),
            (
                "ADGM Constitutional Documents Guide",
                "https://www.adgm.com/legal-framework/guidance-and-policy-statements",
            ),
            (
                "ADGM Jurisdiction Requirements",
                "https://www.adgm.com/legal-framework/regulations",
            ),
        ],
        "missing_clauses": [
            (
                "ADGM Guidance & Templates",
                "https://www.adgm.com/legal-framework/guidance-and-policy-statements",
            ),
            (
                "ADGM Document Templates",
                "https://assets.adgm.com/templates",
            ),
            (
                "ADGM Required Clauses Guide",
                "https://www.adgm.com/legal-framework/guidance-and-policy-statements",
            ),
        ],
        "formatting": [
            (
                "ADGM Registration & Incorporation",
                "https://www.adgm.com/registration-authority/registration-and-incorporation",
            ),
            (
                "ADGM Document Standards",
                "https://www.adgm.com/legal-framework/guidance-and-policy-statements",
            ),
            (
                "ADGM Template Library",
                "https://assets.adgm.com/templates",
            ),
        ],
        "compliance": [
            (
                "ADGM Incorporation Package Rulebook",
                "https://en.adgm.thomsonreuters.com/rulebook/7-company-incorporation-package",
            ),
            (
                "ADGM Compliance Guide",
                "https://www.adgm.com/operating-in-adgm/obligations-of-adgm-registered-entities",
            ),
            (
                "ADGM Regulatory Framework",
                "https://www.adgm.com/legal-framework/regulations",
            ),
        ],
        "ambiguity": [
            (
                "ADGM Legal Writing Guidelines",
                "https://www.adgm.com/legal-framework/guidance-and-policy-statements",
            ),
            (
                "ADGM Template Library",
                "https://assets.adgm.com/templates",
            ),
            (
                "ADGM Plain Language Guide",
                "https://www.adgm.com/legal-framework/guidance-and-policy-statements",
            ),
        ],
        "outdated": [
            (
                "ADGM Guidance & Policy Updates",
                "https://www.adgm.com/legal-framework/guidance-and-policy-statements",
            ),
            (
                "ADGM Regulatory Updates",
                "https://www.adgm.com/legal-framework/regulations",
            ),
            (
                "ADGM Latest Templates",
                "https://assets.adgm.com/templates",
            ),
        ],
    }

    # Collect all relevant links
    all_links = []

    # Add category-specific links
    for category in categories:
        if category in resources_map:
            all_links.extend(resources_map[category])

    # Add document-type specific links
    doc_type_links = get_document_type_specific_links(doc_type)
    if doc_type_links:
        all_links.extend(doc_type_links)

    # Add core reference only if not already included
    core_ref = (
        "ADGM Companies Regulations 2020",
        "https://en.adgm.thomsonreuters.com/rulebook/1-companies-regulations-2020",
    )
    if core_ref not in all_links:
        all_links.append(core_ref)

    # Remove duplicates while preserving order
    seen = set()
    unique_links = []
    for link in all_links:
        if link not in seen:
            seen.add(link)
            unique_links.append(link)

    return unique_links


def get_document_type_specific_links(doc_type: str) -> List[Tuple[str, str]]:
    """Get document-type specific helpful resources"""

    doc_type_links = {
        "Articles of Association (AoA)": [
            (
                "ADGM Articles of Association Template",
                "https://assets.adgm.com/templates/articles-of-association",
            ),
            (
                "ADGM Constitutional Documents Guide",
                "https://www.adgm.com/legal-framework/guidance-and-policy-statements",
            ),
        ],
        "Memorandum of Association (MoA/MoU)": [
            (
                "ADGM Memorandum Template",
                "https://assets.adgm.com/templates/memorandum-of-association",
            ),
            (
                "ADGM Incorporation Requirements",
                "https://www.adgm.com/registration-authority/registration-and-incorporation",
            ),
        ],
        "Board Resolution": [
            (
                "ADGM Resolution Templates",
                "https://assets.adgm.com/templates/resolutions",
            ),
            (
                "ADGM Corporate Governance Guide",
                "https://www.adgm.com/legal-framework/guidance-and-policy-statements",
            ),
        ],
        "Employment Contract": [
            (
                "ADGM Standard Employment Contract (2024)",
                "https://assets.adgm.com/download/assets/ADGM+Standard+Employment+Contract+Template+-+ER+2024+(Feb+2025).docx/ee14b252edbe11efa63b12b3a30e5e3a",
            ),
            (
                "ADGM Employment Regulations",
                "https://en.adgm.thomsonreuters.com/rulebook/employment-regulations",
            ),
        ],
        "UBO Declaration": [
            (
                "ADGM UBO Declaration Form",
                "https://assets.adgm.com/templates/ubo-declaration",
            ),
            (
                "ADGM Beneficial Ownership Guide",
                "https://www.adgm.com/legal-framework/guidance-and-policy-statements",
            ),
        ],
        "Shareholder Agreement": [
            (
                "ADGM Shareholder Agreement Template",
                "https://assets.adgm.com/templates/shareholder-agreement",
            ),
            (
                "ADGM Corporate Governance Standards",
                "https://www.adgm.com/legal-framework/guidance-and-policy-statements",
            ),
        ],
    }

    return doc_type_links.get(doc_type, [])


def insert_comments(
    docx_file, flagged_sections: List[Dict[str, Any]], doc_type: str = "Unknown"
) -> Document:
    """Enhanced comment insertion with better document handling"""
    try:
        doc = Document(docx_file)

        # Add document header with analysis info
        header_para = doc.paragraphs[0].insert_paragraph_before()
        header_run = header_para.add_run("ADGM COMPLIANCE REVIEW - AUTOMATED ANALYSIS")
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(0, 0, 139)  # Dark Blue
        header_run.font.size = Pt(14)
        header_para.add_run(
            f" | Analysis Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}"
        ).font.size = Pt(10)
        header_para.paragraph_format.space_after = Pt(6)

        # Filter out comprehensive error results and handle them separately
        comprehensive_errors = [
            f for f in flagged_sections if f.get("is_comprehensive_error")
        ]
        regular_issues = [
            f for f in flagged_sections if not f.get("is_comprehensive_error")
        ]

        # Insert summary of issues
        if regular_issues or comprehensive_errors:
            summary_para = doc.paragraphs[0].insert_paragraph_before()

            # Handle comprehensive errors
            if comprehensive_errors:
                error_para = summary_para.add_run("⚠️ TECHNICAL NOTICE: ")
                error_para.font.bold = True
                error_para.font.color.rgb = RGBColor(255, 0, 0)  # Red
                error_para.font.size = Pt(12)

                # Show only one consolidated message
                first_error = comprehensive_errors[0]
                summary_para.add_run(
                    f"{first_error.get('red_flag', 'Technical issue occurred')} "
                )
                summary_para.add_run(
                    f"({first_error.get('suggestion', 'Please retry')})"
                )
                summary_para.add_run().add_break()

                summary_para.paragraph_format.space_after = Pt(6)
                return doc  # Don't proceed with regular analysis if there are comprehensive errors

            # Regular issues summary
            high_count = sum(1 for f in regular_issues if f.get("severity") == "High")
            medium_count = sum(
                1 for f in regular_issues if f.get("severity") == "Medium"
            )
            low_count = sum(1 for f in regular_issues if f.get("severity") == "Low")

            # Summary heading
            summary_heading = summary_para.add_run("AI SUMMARY")
            summary_heading.font.bold = True
            summary_heading.font.color.rgb = RGBColor(139, 0, 0)  # Dark Red
            summary_heading.font.size = Pt(12)
            summary_para.add_run(": ")

            # Counts line
            counts_run = summary_para.add_run(
                f"Issues found: {len(regular_issues)}  (High: {high_count}, Medium: {medium_count}, Low: {low_count})"
            )
            counts_run.font.size = Pt(10)
            summary_para.add_run().add_break()

            # Category breakdown
            category_counts: Dict[str, int] = {}
            for f in regular_issues:
                cat = f.get("category", "other")
                category_counts[cat] = category_counts.get(cat, 0) + 1
            if category_counts:
                summary_para.add_run("Categories: ")
                cats_line = ", ".join(f"{k} ({v})" for k, v in category_counts.items())
                cat_run = summary_para.add_run(cats_line)
                cat_run.font.size = Pt(10)
                summary_para.add_run().add_break()

            # Helpful resources (dynamic by categories)
            # Use dynamic link generation based on actual issues found
            all_links = generate_dynamic_helpful_links(regular_issues, doc_type)

            if all_links:
                summary_para.add_run("Helpful resources:")
                summary_para.add_run().add_break()
                # Add each as a bullet-style line with plain text URLs
                for title, url in all_links:
                    summary_para.add_run("- ")
                    # Simple text with URL instead of embedded hyperlink
                    link_run = summary_para.add_run(f"{title}: {url}")
                    link_run.font.color.rgb = RGBColor(0, 102, 204)  # Blue color
                    link_run.font.underline = True
                    summary_para.add_run().add_break()

            summary_para.paragraph_format.space_after = Pt(6)

        # Insert comments at relevant paragraphs for regular issues only
        for flag in regular_issues:
            idx = flag.get("section_index", flag.get("index", 0))
            if idx < len(doc.paragraphs):
                add_comment_to_paragraph(
                    doc.paragraphs[idx],
                    flag.get("red_flag", ""),
                    flag.get("severity", "Medium"),
                    flag.get("law_citation", ""),
                    flag.get("suggestion", ""),
                )

        # Add footer with disclaimer
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run(
            "\n--- END OF ADGM AUTOMATED REVIEW ---\n"
            "DISCLAIMER: This automated review is for guidance only. "
            "Please consult qualified ADGM legal counsel for final compliance verification."
        )
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(105, 105, 105)  # Dim Gray
        footer_run.font.size = Pt(9)
        footer_para.paragraph_format.space_before = Pt(6)

        return doc

    except Exception as e:
        logger.error(f"Error inserting comments: {e}")
        # Return original document if processing fails
        return Document(docx_file)


def save_docx(doc: Document, output_path: str) -> bool:
    """Enhanced document saving with error handling and validation"""
    try:
        # Ensure output directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)

        # Save with error handling
        doc.save(output_path)

        # Validate saved file
        if Path(output_path).exists() and Path(output_path).stat().st_size > 0:
            logger.info(f"Document successfully saved to {output_path}")
            return True
        else:
            logger.error(f"Failed to save document to {output_path}")
            return False

    except Exception as e:
        logger.error(f"Error saving document to {output_path}: {e}")
        return False


def build_red_flag_prompt(
    section_text: str, doc_type: str, context: Dict[str, Any] = None
) -> str:
    """Enhanced prompt building with context-aware analysis"""
    base_prompt = f"""
    You are an expert ADGM legal compliance analyst. Analyze the following section from a {doc_type} for ADGM compliance.
    
    DOCUMENT TYPE: {doc_type}
    SECTION CONTENT: "{section_text}"
    """

    if context:
        base_prompt += f"\nADDITIONAL CONTEXT: {json.dumps(context, indent=2)}"

    base_prompt += """
    
    COMPLIANCE CHECKS REQUIRED:
    1. JURISDICTION: Must reference ADGM Courts, not UAE Federal/Dubai/Abu Dhabi courts
    2. REQUIRED CLAUSES: Check for mandatory sections per ADGM regulations
    3. FORMATTING: Proper structure, signatures, dates as required
    4. TEMPLATE COMPLIANCE: Adherence to official ADGM templates
    5. LANGUAGE CLARITY: Avoid ambiguous or non-binding language
    6. CURRENT REGULATIONS: Ensure references are to current ADGM laws
    
    RESPONSE FORMAT (JSON):
    {
        "red_flag": "Detailed description of issue found, or null if compliant",
        "law_citation": "Exact ADGM regulation citation (e.g., 'ADGM Companies Regulations 2020, Article X')",
        "suggestion": "Specific compliant alternative wording or action",
        "severity": "High/Medium/Low based on legal and business impact",
        "category": "jurisdiction/missing_clauses/formatting/compliance/ambiguity/outdated",
        "confidence": "High/Medium/Low based on analysis certainty",
        "compliant_clause": "Suggested replacement clause text if applicable"
    }
    
    Focus on ADGM-specific requirements and provide precise citations.
    """

    return base_prompt


def extract_document_metadata(docx_file) -> Dict[str, Any]:
    """Enhanced metadata extraction with comprehensive document analysis"""
    try:
        doc = Document(docx_file)

        metadata = {
            "paragraph_count": len(doc.paragraphs),
            "section_count": len(doc.sections),
            "table_count": len(doc.tables),
            "headers": [],
            "signatures": [],
            "dates": [],
            "defined_terms": [],
            "financial_amounts": [],
            "jurisdictions": [],
            "document_length": 0,
            "structure_analysis": {},
        }

        # Analyze paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            metadata["document_length"] += len(text)

            # Extract headers
            if (
                (para.style and para.style.name.startswith("Heading"))
                or (len(text) < 100 and text.isupper())
                or text.endswith(":")
            ):
                metadata["headers"].append(text)

            # Extract signatures
            for pattern in SIGNATURE_PATTERNS:
                if re.search(pattern, text.lower()):
                    metadata["signatures"].append(text)
                    break

            # Extract dates
            for pattern in DATE_PATTERNS:
                dates_found = re.findall(pattern, text)
                metadata["dates"].extend(dates_found)

            # Extract defined terms (quoted text)
            defined_terms = re.findall(r'"([^"]+)"', text)
            metadata["defined_terms"].extend(defined_terms)

            # Extract financial amounts
            amounts = re.findall(r"[\$£€]\s*[\d,]+(?:\.\d{2})?", text)
            amounts.extend(
                re.findall(r"\b\d+(?:\.\d{2})?\s*(?:USD|AED|EUR|GBP)\b", text)
            )
            metadata["financial_amounts"].extend(amounts)

            # Extract jurisdiction references
            jurisdiction_patterns = [
                r"ADGM\s+Court",
                r"Abu\s+Dhabi\s+Global\s+Market",
                r"UAE\s+Federal\s+Court",
                r"Dubai\s+Court",
                r"Abu\s+Dhabi\s+Court",
            ]
            for pattern in jurisdiction_patterns:
                if re.search(pattern, text, re.I):
                    metadata["jurisdictions"].append(
                        re.search(pattern, text, re.I).group()
                    )

        # Structure analysis
        metadata["structure_analysis"] = {
            "has_table_of_contents": any(
                "content" in h.lower() for h in metadata["headers"]
            ),
            "has_signatures": len(metadata["signatures"]) > 0,
            "has_dates": len(metadata["dates"]) > 0,
            "has_defined_terms": len(metadata["defined_terms"]) > 0,
            "has_financial_info": len(metadata["financial_amounts"]) > 0,
            "adgm_jurisdiction": any(
                "adgm" in j.lower() for j in metadata["jurisdictions"]
            ),
            "incorrect_jurisdiction": any(
                any(
                    court in j.lower()
                    for court in ["uae federal", "dubai court", "abu dhabi court"]
                )
                for j in metadata["jurisdictions"]
            ),
        }

        # Clean up duplicates
        metadata["headers"] = list(set(metadata["headers"]))
        metadata["signatures"] = list(set(metadata["signatures"]))
        metadata["dates"] = list(set(metadata["dates"]))
        metadata["defined_terms"] = list(set(metadata["defined_terms"]))

        return metadata

    except Exception as e:
        logger.error(f"Error extracting document metadata: {e}")
        return {"error": str(e)}


def analyze_document_completeness(
    sections: List[Dict[str, Any]], doc_type: str
) -> Dict[str, Any]:
    """Analyze document completeness against ADGM requirements"""
    required_elements = REQUIRED_SECTIONS.get(doc_type, [])

    analysis = {
        "doc_type": doc_type,
        "total_sections": len(sections),
        "required_elements": required_elements,
        "found_elements": [],
        "missing_elements": [],
        "completeness_score": 0.0,
        "recommendations": [],
    }

    if not required_elements:
        analysis["completeness_score"] = 100.0
        analysis["recommendations"].append(
            "No specific structure requirements defined for this document type"
        )
        return analysis

    # Check for required elements
    content_text = " ".join([s.get("text", "").lower() for s in sections])

    for element in required_elements:
        if element.lower() in content_text:
            analysis["found_elements"].append(element)
        else:
            analysis["missing_elements"].append(element)

    # Calculate completeness score
    if required_elements:
        analysis["completeness_score"] = (
            len(analysis["found_elements"]) / len(required_elements)
        ) * 100

    # Generate recommendations
    if analysis["missing_elements"]:
        analysis["recommendations"].append(
            f"Add missing sections: {', '.join(analysis['missing_elements'][:3])}"
        )

    if analysis["completeness_score"] >= 80:
        analysis["recommendations"].append("Document structure is largely complete")
    elif analysis["completeness_score"] >= 60:
        analysis["recommendations"].append(
            "Document needs minor structural improvements"
        )
    else:
        analysis["recommendations"].append(
            "Document requires significant structural review"
        )

    return analysis


def create_compliance_report(
    analysis_results: List[Dict[str, Any]], doc_info: Dict[str, Any]
) -> Dict[str, Any]:
    """Create comprehensive compliance report"""
    report = {
        "document_info": doc_info,
        "analysis_timestamp": datetime.datetime.now().isoformat(),
        "total_issues": len(analysis_results),
        "severity_breakdown": {"High": 0, "Medium": 0, "Low": 0},
        "category_breakdown": {},
        "compliance_score": 100.0,
        "issues": analysis_results,
        "recommendations": [],
        "next_steps": [],
    }

    # Calculate severity and category breakdowns
    for issue in analysis_results:
        severity = issue.get("severity", "Medium")
        category = issue.get("category", "general")

        report["severity_breakdown"][severity] += 1
        report["category_breakdown"][category] = (
            report["category_breakdown"].get(category, 0) + 1
        )

    # Calculate compliance score
    if analysis_results:
        high_weight = report["severity_breakdown"]["High"] * 10
        medium_weight = report["severity_breakdown"]["Medium"] * 5
        low_weight = report["severity_breakdown"]["Low"] * 2
        total_weight = high_weight + medium_weight + low_weight

        # Score decreases with weighted issues
        max_possible = len(analysis_results) * 10
        report["compliance_score"] = max(0, 100 - (total_weight / max_possible * 100))

    # Generate recommendations
    if report["severity_breakdown"]["High"] > 0:
        report["recommendations"].append(
            "Address high-priority compliance issues immediately"
        )
        report["next_steps"].append(
            "Review jurisdiction clauses and critical missing elements"
        )

    if report["severity_breakdown"]["Medium"] > 0:
        report["recommendations"].append(
            "Plan to resolve medium-priority issues within 2-3 business days"
        )

    if report["total_issues"] == 0:
        report["recommendations"].append(
            "Document appears compliant with ADGM requirements"
        )
        report["next_steps"].append(
            "Consider final review with qualified ADGM legal counsel"
        )

    return report


# Export all functions
__all__ = [
    "extract_text_sections",
    "detect_document_type",
    "validate_document_structure",
    "add_comment_to_paragraph",
    "insert_comments",
    "save_docx",
    "build_red_flag_prompt",
    "extract_document_metadata",
    "analyze_document_completeness",
    "create_compliance_report",
]
